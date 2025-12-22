from docx import Document
from datetime import datetime
import io
import os

class DocumentService:
    def generate_from_template(self, template_path: str, context: dict) -> io.BytesIO:
        """
        Generates a DOCX file by replacing placeholders in a template.
        template_path: Absolute path to .docx template via python-docx.
        context: Dictionay of {placeholder: value}.
        Returns: BytesIO object of the generated file.
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)

        # 1. Replace in Paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, context)

        # 2. Replace in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, context)

        # Save to buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def _replace_text_in_paragraph(self, paragraph, context):
        """
        Helper to replace text in a paragraph. 
        Note: This is a simple replacement. Complex formatting might break if run spans are split.
        """
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}" # Expects {{key}}
            if placeholder in paragraph.text:
                # Simple replacement (might lose formatting if placeholder spans multiple runs)
                # For robust replacement, we need to iterate runs, but this suffices for MVP
                paragraph.text = paragraph.text.replace(placeholder, str(value))

document_service = DocumentService()
