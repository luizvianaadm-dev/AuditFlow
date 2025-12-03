from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
from datetime import datetime

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) if level == 0 else RGBColor(0, 174, 239)
    if level == 0:
        run.font.size = Pt(24)
    else:
        run.font.size = Pt(14)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    # Simple HTML-like parsing for bold
    parts = text.split('<b>')
    for i, part in enumerate(parts):
        if '</b>' in part:
            subparts = part.split('</b>')
            # Bold part
            run = p.add_run(subparts[0])
            run.bold = True
            # Normal part
            if len(subparts) > 1:
                p.add_run(subparts[1])
        else:
            # First part or no closing tag
            run = p.add_run(part)
            if bold: run.bold = True
            if italic: run.italic = True

def generate_audit_report_docx(engagement, analysis_results, mistatement_summary=None):
    doc = Document()

    # 1. Header / Cover
    add_heading(doc, "Relatório de Auditoria Digital", level=0)

    # Engagement Details Table
    table = doc.add_table(rows=4, cols=2)
    data = [
        ('Cliente:', engagement.client.name),
        ('Auditoria:', engagement.name),
        ('Ano Base:', str(engagement.year)),
        ('Data do Relatório:', datetime.now().strftime("%d/%m/%Y"))
    ]
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    doc.add_paragraph() # Spacer

    # 2. Executive Summary
    add_heading(doc, "Resumo Executivo", level=1)
    summary_text = f"Este documento apresenta os resultados da auditoria automatizada realizada para o cliente <b>{engagement.client.name}</b>. Foram analisadas {len(engagement.transactions)} transações financeiras utilizando testes estatísticos e forenses conforme as normas NBC TA 240 e 520."
    add_paragraph(doc, summary_text)
    doc.add_paragraph()

    # 3. Analysis Results
    if not analysis_results:
        doc.add_paragraph("Nenhuma análise foi registrada para este trabalho.")
    else:
        for analysis in analysis_results:
            test_name = "Lei de Benford" if analysis.test_type == 'benford' else "Pagamentos Duplicados"
            if analysis.test_type == 'materiality': test_name = "Cálculo de Materialidade"

            add_heading(doc, f"Resultado: {test_name}", level=1)
            add_paragraph(doc, f"Executado em: {analysis.executed_at.strftime('%d/%m/%Y %H:%M')}", italic=True)

            if analysis.test_type == 'benford':
                anomalies = analysis.result.get('anomalies', [])
                if anomalies:
                    add_paragraph(doc, f"<b>{len(anomalies)} Anomalias Estatísticas Detectadas</b>")
                    doc.add_paragraph("Os seguintes dígitos apresentaram desvio significativo (>5%) da frequência esperada:")

                    # Table
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Dígito'
                    hdr_cells[1].text = 'Esperado'
                    hdr_cells[2].text = 'Observado'
                    hdr_cells[3].text = 'Desvio'

                    details = analysis.result.get('details', [])
                    for d in details:
                        if d['is_anomaly']:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(d['digit'])
                            row_cells[1].text = f"{d['expected']*100:.1f}%"
                            row_cells[2].text = f"{d['observed']*100:.1f}%"
                            row_cells[3].text = f"{d['deviation']*100:.1f}%"
                else:
                    doc.add_paragraph("Nenhuma anomalia estatística detectada. A distribuição segue o padrão esperado.")

            elif analysis.test_type == 'duplicates':
                groups = analysis.result.get('duplicates', [])
                if groups:
                    add_paragraph(doc, f"<b>{len(groups)} Grupos de Pagamentos Suspeitos</b>")

                    for i, group in enumerate(groups[:10]):
                        p = doc.add_paragraph(f"Grupo {i+1} - Valor: R$ {group['amount']:.2f} (Similaridade: {group['similarity_score']}%)")
                        for tx in group['transactions']:
                            doc.add_paragraph(f"- {tx['vendor']} (ID: {tx['id']})", style='List Bullet')

                    if len(groups) > 10:
                        add_paragraph(doc, f"...e mais {len(groups)-10} grupos.", italic=True)
                else:
                    doc.add_paragraph("Nenhum pagamento duplicado encontrado com os critérios atuais.")

            elif analysis.test_type == 'materiality':
                res = analysis.result
                add_paragraph(doc, f"<b>Materialidade Global: R$ {res.get('global_materiality', 0):,.2f}</b>")
                doc.add_paragraph(f"Base de Cálculo: {res.get('benchmark')} (R$ {res.get('benchmark_value', 0):,.2f})")
                doc.add_paragraph(f"Percentual Aplicado: {res.get('percentage_global')}%")
                doc.add_paragraph(f"Materialidade de Performance: R$ {res.get('performance_materiality', 0):,.2f}")

            doc.add_paragraph()

    # 4. Summary of Mistatements
    if mistatement_summary and mistatement_summary.get('items'):
        add_heading(doc, "Sumário de Erros Não Corrigidos", level=1)
        doc.add_paragraph("Abaixo listamos as divergências identificadas durante os trabalhos de auditoria:")

        items = mistatement_summary['items']
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Descrição'
        hdr_cells[1].text = 'Tipo'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Valor (R$)'

        for m in items:
            row_cells = table.add_row().cells
            row_cells[0].text = m.description
            row_cells[1].text = m.type.capitalize()
            row_cells[2].text = "Ajustado" if m.status == 'adjusted' else "Não Ajustado"
            row_cells[3].text = f"{m.amount_divergence:,.2f}"

        doc.add_paragraph()
        add_paragraph(doc, f"<b>Total Ajustado: R$ {mistatement_summary['total_adjusted']:,.2f}</b>")
        add_paragraph(doc, f"<b>Total Não Ajustado: R$ {mistatement_summary['total_unadjusted']:,.2f}</b>")

        # Conclusion
        materiality_val = 0
        for res in analysis_results:
            if res.test_type == 'materiality':
                materiality_val = res.result.get('global_materiality', 0)
                break

        if materiality_val > 0:
            doc.add_paragraph(f"Materialidade Global: R$ {materiality_val:,.2f}")

            if mistatement_summary['total_unadjusted'] > materiality_val:
                conclusion = "CONCLUSÃO: O total de erros não ajustados excede a materialidade global. As demonstrações contábeis podem conter distorções relevantes."
                p = doc.add_paragraph()
                run = p.add_run(conclusion)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                conclusion = "CONCLUSÃO: O total de erros não ajustados é inferior à materialidade global. As distorções não são consideradas relevantes individualmente ou em conjunto."
                p = doc.add_paragraph()
                run = p.add_run(conclusion)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)

    # Footer
    doc.add_paragraph()
    add_paragraph(doc, "Gerado automaticamente por AuditFlow", italic=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
